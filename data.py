import sqlite3  
import flet as ft 
import pandas as pd
from fpdf import FPDF
import os
import platform
import subprocess
from datetime import datetime
from zoneinfo import ZoneInfo
from docx import Document 
from datetime import datetime
import sys
import subprocess

db_path = "inventory.db"

def abrir_pdf(caminho_pdf):
    if sys.platform.startswith('darwin'):
        subprocess.call(('open', caminho_pdf))
    elif os.name == 'nt':
        os.startfile(caminho_pdf)
    elif os.name == 'posix':
        subprocess.call(('xdg-open', caminho_pdf))

def get_connection():
    return sqlite3.connect(db_path)

def ensure_tables_exist():
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS equipamento (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                data_entrada TEXT,
                equipamento TEXT,
                marca TEXT,
                quantidade INTEGER,
                validade TEXT,
                observacoes TEXT,
                data_saida TEXT
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS patrimonio (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                data_entrada TEXT,
                equipamento TEXT,
                marca TEXT,
                quantidade INTEGER,
                validade TEXT,
                observacoes TEXT,
                data_saida TEXT
            )
        ''')
        conn.commit()

def format_date_input(field, page):
    valor = field.value.strip()
    if len(valor) == 8 and "/" not in valor:
        valor = f"{valor[:2]}/{valor[2:4]}/{valor[4:]}"
        field.value = valor
        page.update()
    partes = field.value.split("/")
    if len(partes) == 3 and len(partes[2]) > 4:
        partes[2] = partes[2][:4]
        field.value = "/".join(partes)
        page.update()
    try:
        datetime.strptime(field.value, "%d/%m/%Y")
    except ValueError:
        show_popup(page, "Erro", "Data inválida! Use DD/MM/AAAA.")

def process_item(tipo, equipamento_input, marca_input, quantidade_input, validade_input, observacoes_input, saida_input, page):
    try:
        quantidade = int(quantidade_input.value.strip())
    except ValueError:
        show_popup(page, "Erro", "Quantidade inválida!")
        return
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(f'''
                INSERT INTO {tipo} (data_entrada, equipamento, marca, quantidade, validade, observacoes, data_saida)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                datetime.now().strftime("%d/%m/%Y"),
                equipamento_input.value.strip(),
                marca_input.value.strip(),
                quantidade,
                validade_input.value.strip(),
                observacoes_input.value.strip(),
                saida_input.value.strip()
            ))
            conn.commit()
        equipamento_input.value = ""
        marca_input.value = ""
        quantidade_input.value = ""
        validade_input.value = ""
        observacoes_input.value = ""
        saida_input.value = ""
        show_popup(page, "Sucesso", f"{tipo.capitalize()} adicionado com sucesso!")
    except Exception as e:
        show_popup(page, "Erro", f"Erro ao adicionar: {e}")

def load_items_for_report(tipo):
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT id, data_entrada, equipamento, marca, quantidade, validade, observacoes, data_saida FROM {tipo}")
        return [list(item) for item in cursor.fetchall()]

def save_edited_items(tipo, items):
    with get_connection() as conn:
        cursor = conn.cursor()
        for item in items:
            if item[0]:
                cursor.execute(f'''
                    UPDATE {tipo}
                    SET data_entrada = ?, equipamento = ?, marca = ?, quantidade = ?,
                        validade = ?, observacoes = ?, data_saida = ?
                    WHERE id = ?
                ''', (
                    item[1], item[2], item[3], int(item[4]),
                    item[5], item[6], item[7], item[0]
                ))
            else:
                cursor.execute(f'''
                    INSERT INTO {tipo} (data_entrada, equipamento, marca, quantidade,
                                        validade, observacoes, data_saida)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    item[1], item[2], item[3], int(item[4]),
                    item[5], item[6], item[7]
                ))
        conn.commit()

def delete_item(tipo, item_id):
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(f"DELETE FROM {tipo} WHERE id = ?", (item_id,)) 
        conn.commit()

def show_popup(page: ft.Page, titulo: str, mensagem: str):
    dialog = ft.AlertDialog(
        title=ft.Text(titulo),
        content=ft.Text(mensagem),
        actions=[ft.TextButton("OK", on_click=lambda e: fechar_dialog(dialog, page))],
        actions_alignment=ft.MainAxisAlignment.END,
    )
    page.dialog = dialog
    dialog.open = True
    page.update()

def fechar_dialog(dialog, page):
    dialog.open = False
    page.update()

def imprimir_pdf(path_pdf):
    sistema = platform.system()
    try:
        if sistema == "Windows":
            os.startfile(path_pdf, "print")
        elif sistema == "Darwin":
            subprocess.run(["lp", path_pdf])
        elif sistema == "Linux":
            subprocess.run(["lp", path_pdf])
    except Exception as e:
        print(f"Erro ao imprimir: {e}")

RelatoriosDiario = os.path.join(os.path.expanduser('~'), 'Desktop')
PASTA_RELATORIOS_DIARIO = os.path.join(RelatoriosDiario, "RelatoriosDiários") 
area_de_trabalho = os.path.join(os.path.expanduser('~'), 'Desktop') 
PASTA_RELATORIOS = os.path.join(area_de_trabalho, "OutrosRelatorios")


def get_nome_relatorio(tipo, extensao):
    data_str = datetime.now().strftime("%d_%m_%Y_%H_%M")
    return f"{tipo}_{data_str}.{extensao}"

def export_to_excel(tipo, items):
    if not os.path.exists(PASTA_RELATORIOS):
        os.makedirs(PASTA_RELATORIOS)
    nome_arquivo = get_nome_relatorio(tipo, "xlsx")
    caminho = os.path.join(PASTA_RELATORIOS, nome_arquivo)
    dados_sem_id = [item[1:] for item in items]
    df = pd.DataFrame(dados_sem_id, columns=[
        "Data Entrada", "Equipamento", "Marca", "Quantidade", "Validade", "Observacoes", "Data Saida"
    ])
    df.to_excel(caminho, index=False)
    return caminho

def export_to_pdf(tipo, items):
    if not os.path.exists(PASTA_RELATORIOS):
        os.makedirs(PASTA_RELATORIOS)
    nome_arquivo = get_nome_relatorio(tipo, "pdf")
    caminho = os.path.join(PASTA_RELATORIOS, nome_arquivo)

    pdf = FPDF(orientation="L")
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    data_rel = datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%d/%m/%Y")
    pdf.cell(0, 5, f"Relatório gerado em: {data_rel}", ln=1)
    pdf.ln(5)

    headers = ["Data Entrada", "Equipamento", "Marca", "Qtd", "Validade", "Obs", "Data Saída"]
    th = pdf.font_size + 3
    largura_util = pdf.w - 2 * pdf.l_margin
    col_pesos = [1, 2, 1.5, 0.7, 1.2, 2, 1.2]
    col_widths = [largura_util * (peso / sum(col_pesos)) for peso in col_pesos]

    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], th, header, border=1)
    pdf.ln(th)

    for row in items:
        y_start = pdf.get_y()
        max_y = y_start
        cell_texts = row[1:]
        x_start = pdf.get_x()
        y_start = pdf.get_y()

        for i, text in enumerate(cell_texts):
            texto = str(text)
            x = pdf.get_x()
            y = pdf.get_y()

            if i in [1, 5]:
                pdf.multi_cell(col_widths[i], th, texto, border=1)
                max_y = max(max_y, pdf.get_y())
                pdf.set_xy(x + col_widths[i], y)
            else:
                pdf.cell(col_widths[i], th, texto, border=1)
        pdf.ln(max_y - y_start)

    pdf.ln(20)
    pdf.set_x(pdf.l_margin)
    pdf.cell(pdf.w - 2 * pdf.l_margin, 0, '', 'T', ln=1)
    pdf.ln(5)
    pdf.cell(0, 5, 'Bombeiro Líder', 0, 1, 'C')

    pdf.output(caminho)
    return caminho

def export_to_word(tipo, items):
    if not os.path.exists(PASTA_RELATORIOS):
        os.makedirs(PASTA_RELATORIOS)

    nome_arquivo = get_nome_relatorio(tipo, "docx")
    caminho = os.path.join(PASTA_RELATORIOS, nome_arquivo)

    doc = Document()
    doc.add_heading(f'Relatório de {tipo.capitalize()}', 0)

    headers = ["Data Entrada", "Equipamento", "Marca", "Quantidade", "Validade", "Observações", "Data Saída"]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for item in items:
        row_cells = table.add_row().cells
        for i in range(len(headers)):
            row_cells[i].text = str(item[i+1])  # pula o ID

    doc.save(caminho)
    return caminho

def listar_pdfs():
    """
    Retorna a lista de nomes de arquivos .pdf dentro da pasta de relatórios.
    """
    if not os.path.exists(PASTA_RELATORIOS_DIARIO):
        return []
    return [
        nome for nome in os.listdir(PASTA_RELATORIOS_DIARIO)
        if nome.lower().endswith(".pdf")
    ]

def ensure_snapshot_tables_exist():
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS snapshot_equipamento (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                snapshot_date TEXT,
                snapshot_type TEXT,
                equipamento_id INTEGER,
                data_entrada TEXT, equipamento TEXT, marca TEXT,
                quantidade INTEGER, validade TEXT,
                observacoes TEXT, data_saida TEXT
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS snapshot_patrimonio (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                snapshot_date TEXT,
                snapshot_type TEXT,
                equipamento_id INTEGER,
                data_entrada TEXT, equipamento TEXT, marca TEXT,
                quantidade INTEGER, validade TEXT,
                observacoes TEXT, data_saida TEXT
            )
        ''')
        conn.commit()

def registrar_snapshot(tipo: str, snapshot_type: str):
    """
    Grava todos os itens atuais de ‘tipo’ (equipamento|patrimonio) 
    na tabela snapshot_{tipo} com snapshot_type = 'inicio' ou 'fim'.
    """
    items = load_items_for_report(tipo)
    agora = datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%Y-%m-%d %H:%M:%S")
    tabela = f"snapshot_{tipo}"
    with get_connection() as conn:
        cursor = conn.cursor()
        for row in items:
            equipamento_id = row[0]
            dados = row[1:]  
            cursor.execute(f'''
                INSERT INTO {tabela}(
                  snapshot_date, snapshot_type, equipamento_id,
                  data_entrada, equipamento, marca, quantidade,
                  validade, observacoes, data_saida
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', [agora, snapshot_type, equipamento_id] + dados)
        conn.commit()

def get_snapshots(tipo: str, date_str: str, snapshot_type: str):
    tabela = f"snapshot_{tipo}"
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(f'''
            SELECT equipamento_id, data_entrada, equipamento, marca,
                   quantidade, validade, observacoes, data_saida
            FROM {tabela}
            WHERE DATE(snapshot_date) = ? AND snapshot_type = ?
        ''', (date_str, snapshot_type))
        return [list(r) for r in cursor.fetchall()]

def calculate_diffs(tipo: str, date_str: str):
    inicio = get_snapshots(tipo, date_str, "inicio")
    fim    = get_snapshots(tipo, date_str, "fim")
    d0 = {r[0]: r for r in inicio}
    d1 = {r[0]: r for r in fim}
    added   = [r for k, r in d1.items() if k not in d0]
    removed = [r for k, r in d0.items() if k not in d1]
    altered = []
    for k in d0.keys() & d1.keys():
        if d0[k][1:] != d1[k][1:]:
            altered.append((d0[k], d1[k]))
    return added, removed, altered

def export_relatorio_diario(date_str_ddmmyyyy: str):
    from datetime import datetime
    from fpdf import FPDF
    import os

    # Converte data e prepara caminho de saída
    dt = datetime.strptime(date_str_ddmmyyyy, "%d/%m/%Y").date()
    iso_date = dt.isoformat()
    desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
    out_dir = os.path.join(desktop, "RelatoriosDiários")
    os.makedirs(out_dir, exist_ok=True)
    arquivo = os.path.join(out_dir, f"diario_{dt.strftime('%d_%m_%Y')}.pdf")

    # Inicializa PDF
    pdf = FPDF(orientation="P")
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, f"Relatório Diário - {date_str_ddmmyyyy}", ln=1, align="C")
    pdf.ln(4)

    # Configurações de coluna
    pesos   = [1.4, 2.5, 2, 1.4, 1.5, 3, 1.5]
    largura = pdf.w - 2*pdf.l_margin
    col_ws  = [largura * (p / sum(pesos)) for p in pesos]
    th      = pdf.font_size + 3

    for tipo in ("equipamento", "patrimonio"):
        label_item = "Equipamento" if tipo == "equipamento" else "Patrimônio"
        headers = ["Data Entrada", label_item, "Marca", "Quantidade",
                   "Validade", "Observações", "Data Saída"]

        for fase in ("inicio", "fim"):
            # Título de fase
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, f"{label_item} - {fase.capitalize()}", ln=1)
            pdf.ln(1)

            # Cabeçalhos da tabela
            pdf.set_font("Arial", "B", 9)
            for w, h in zip(col_ws, headers):
                pdf.cell(w, th, h, border=1, align="C")
            pdf.ln(th)

            # Linhas de dados
            rows = get_snapshots(tipo, iso_date, fase)
            pdf.set_font("Arial", size=9)
            for row in rows:
                vals = [str(c) for c in row[1:]]  # ignora ID

                # 1) calcula altura da linha
                line_counts = []
                for i, txt in enumerate(vals):
                    if headers[i] == "Observações":
                        lines = pdf.multi_cell(col_ws[i], th, txt, split_only=True)
                        line_counts.append(len(lines))
                    else:
                        line_counts.append(1)
                max_l = max(line_counts)
                row_h = th * max_l

                # 2) verifica e faz quebra de página se necessário
                if pdf.get_y() + row_h > pdf.page_break_trigger:
                    pdf.add_page()
                    # (Opcional) repetir cabeçalho neste ponto

                # 3) imprime cada célula com altura uniforme
                x0, y0 = pdf.get_x(), pdf.get_y()
                for i, txt in enumerate(vals):
                    w = col_ws[i]
                    if headers[i] == "Observações":
                        pdf.multi_cell(w, th, txt, border=1)
                    else:
                        pdf.cell(w, row_h, txt, border=1)
                    pdf.set_xy(x0 + sum(col_ws[:i+1]), y0)

                # 4) avança na vertical
                pdf.ln(row_h)

            pdf.ln(6)

        # Seção de diferenças
        added, removed, altered = calculate_diffs(tipo, dt.strftime("%Y-%m-%d"))
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 8, f"Diferenças em {label_item}:", ln=1)
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 6, "- Itens adicionados:", ln=1)
        if added:
            for r in added:
                pdf.cell(0, 6, f"   - {r[2]} (quantidade: {r[4]})", ln=1)
        else:
            pdf.cell(0, 6, "   - Nenhum item adicionado.", ln=1)
        pdf.cell(0, 6, "- Itens retirados:", ln=1)
        if removed:
            for r in removed:
                pdf.cell(0, 6, f"   - {r[2]} (quantidade: {r[4]})", ln=1)
        else:
            pdf.cell(0, 6, "   - Nenhum item retirado.", ln=1)
        pdf.cell(0, 6, "- Itens alterados:", ln=1)
        if altered:
            for before, after in altered:
                nome = after[2]; qtd0 = before[4]; qtd1 = after[4]
                pdf.cell(0, 6, f"   - {nome} (quantidade: {qtd0} → {qtd1})", ln=1)
        else:
            pdf.cell(0, 6, "   - Nenhum item alterado.", ln=1)
        pdf.ln(8)

    # Rodapé
    pdf.set_x(pdf.l_margin)
    pdf.cell(largura, 0, "", "T", ln=1)
    pdf.ln(3)
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 6, "Bombeiro Líder", ln=1, align="C")

    pdf.output(arquivo)
    return arquivo


ensure_snapshot_tables_exist()

ensure_tables_exist()


